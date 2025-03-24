#!/usr/bin/env python3
import sys
import os
import datetime
import google.generativeai as genai
import base64
from google import genai
from google.genai import types
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import whisper

def generate_report(transcription_file_path, output_path):
    """
    Genera un report di meeting utilizzando Gemini 2.0 Flash e lo salva in formato DOCX con formattazione migliorata.

    Args:
        transcription_file_path (str): Il percorso del file di testo contenente la trascrizione.
        output_path (str): Il percorso completo dove salvare il file DOCX.
    """
    print("Generazione del report di meeting con Gemini 2.0 Flash...")
    try:
        client = genai.Client(
            api_key=os.environ.get("GEMINI_API_KEY"),
        )

        files = [
            # Make the file available in local system working directory
            client.files.upload(file=transcription_file_path),
        ]
        model = "gemini-2.0-flash"
        contents = [
            types.Content(
                role="user",
                parts=[
                    types.Part.from_uri(
                        file_uri=files[0].uri,
                        mime_type=files[0].mime_type,
                    ),
                    types.Part.from_text(text="""Sulla base della trascrizione della riunione, crea un resoconto approfondito, completo e dettagliato degli argomenti discussi, senza omettere alcun intervento o argomento, non essere troppo sintetico, non banalizzare i concetti e gli argomenti discussi. Utilizza un linguaggio tecnico e professionale. Individua come prima cosa un titolo dell'incontro sulla base degli argomenti discussi. Individua i partecipanti all'incontro ed elencali in apertura, indicando anche la data dell'incontro che puoi ricavare dal titolo del file. Struttura il report in paragrafi suddividendolo per tematiche omogenee. Non citare mai direttamente gli speaker. Al termine del resoconto evidenzia gli appuntamenti futuri, le azioni da mettere in atto e a chi si riferiscono. All'inizio del documento crea un executive summary di quanto discusso, delle decisioni prese e delle cose da fare, in forma di punto elenco. Restituisci il testo in italiano senza eccedere in sinteticità, è fondamentale la completezza delle informazioni desumibili."""),
                ],
            ),
        ]
        generate_content_config = types.GenerateContentConfig(
            temperature=0.2,
            top_p=1.00,
            top_k=40,
            max_output_tokens=60000,
            response_mime_type="text/plain",
        )

        report_text = ""
        for chunk in client.models.generate_content_stream(
            model=model,
            contents=contents,
            config=generate_content_config,
        ):
            report_text += chunk.text

        # Crea un nuovo documento DOCX
        document = Document()

        # Imposta il font predefinito a Calibri Light
        styles = document.styles
        for style in styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                style.font.name = 'Calibri Light'
                style._element.rPr.rFonts.set(qn('w:asciiTheme'), 'Calibri Light')
                style._element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Calibri Light')

        # Estrai il titolo del report (assumendo sia la prima riga non vuota)
        report_title = None
        for line in report_text.splitlines():
            line = line.strip()
            if line:
                report_title = line.lstrip('#').strip() # Rimuovi eventuali ## iniziali e spazi
                document.add_heading(report_title, level=1)
                break

        if report_title:
            # Pulisci il titolo per usarlo come nome file
            report_title_cleaned = re.sub(r'[^\w\s-]', '', report_title).strip()
            report_title_cleaned = report_title_cleaned.replace(' ', '_')

            # Ottieni la data odierna per il nome del file
            today_str_report = datetime.datetime.now().strftime("%d_%m_%Y")

            # Crea il nome del file di output per il report con estensione .docx
            report_filename = f"{report_title_cleaned}_{today_str_report}.docx"
        else:
            # Se non si riesce a estrarre il titolo, usa un nome predefinito con la data
            today_str_report = datetime.datetime.now().strftime("%d_%m_%Y")
            report_filename = f"report_senza_titolo_{today_str_report}.docx"

        report_filepath = os.path.join(output_path, report_filename)

        # Flag per indicare se siamo nella sezione dei partecipanti
        is_participants_section = False

        # Aggiungi il resto del testo al documento con formattazione
        for line in report_text.splitlines():
            line = line.strip()
            if line:
                if line.startswith("###"):
                    document.add_heading(line[3:].strip(), level=2)
                    is_participants_section = False # Le sezioni successive non sono partecipanti
                elif line.startswith("* "):
                    bullet_content = line[2:].strip()
                    bold_title_match = re.match(r'^\*\*[^\*]+\*\*:?$', bullet_content)
                    if bold_title_match:
                        paragraph = document.add_paragraph()
                        parts = re.split(r'(\*\*[^\*]+\*\*)', bullet_content)
                        for part in parts:
                            if part.startswith("**") and part.endswith("**"):
                                paragraph.add_run(part[2:-2]).bold = True
                            else:
                                paragraph.add_run(part)
                    else:
                        paragraph = document.add_paragraph(style='List Bullet')
                        parts = re.split(r'(\*\*[^\*]+\*\*)', bullet_content)
                        for part in parts:
                            if part.startswith("**") and part.endswith("**"):
                                paragraph.add_run(part[2:-2]).bold = True
                            else:
                                paragraph.add_run(part)
                    is_participants_section = False # Le liste non sono partecipanti di default
                elif "**" in line:
                    paragraph = document.add_paragraph()
                    parts = re.split(r'(\*\*[^\*]+\*\*)', line)
                    for part in parts:
                        if part.startswith("**") and part.endswith("**"):
                            paragraph.add_run(part[2:-2]).bold = True
                        else:
                            paragraph.add_run(part)
                    if "Partecipanti:" in line:
                        is_participants_section = True
                elif is_participants_section:
                    document.add_paragraph(line, style='List Bullet')
                elif not line.startswith("##"): # Gestisci le righe che iniziano con ## (es. la data)
                    document.add_paragraph(line.lstrip('#').strip()) # Rimuovi eventuali # iniziali

        # Salva il documento DOCX
        document.save(report_filepath)

        print("Report di meeting generato e salvato in:", report_filepath)

    except Exception as e:
        print(f"Si è verificato un errore durante la generazione del report: {e}")

def main():
    # Verifica che venga passato almeno 1 argomento (il file audio)
    if len(sys.argv) < 2:
        print("Uso: python script_name.py /percorso/del/file_audio")
        sys.exit(1)

    audio_file = sys.argv[1]

    # Carica il modello Whisper (puoi cambiare "large" con "small", "medium", ecc.)
    print("Caricamento del modello Whisper...")
    model = whisper.load_model("large")

    # Avvia la trascrizione
    print("Esecuzione della trascrizione...")
    result = model.transcribe(audio_file)

    # Genera la data e l'ora odierne in formato gg_mm_yyyy_HH_MM
    now = datetime.datetime.now()
    timestamp_str = now.strftime("%d_%m_%Y_%H_%M")

    # Crea il nome del file di output con la data e l'ora
    # Ad esempio: "trascrizione_15_04_2025_10_30.txt"
    output_filename = f"trascrizione_{timestamp_str}.txt"

    # Se vuoi salvare il file nella stessa cartella dello script, potresti usare:
    #script_dir = os.path.dirname(__file__)
    #output_path = os.path.join(script_dir, output_filename)

    # Oppure, se vuoi un percorso specifico, usalo direttamente (commenta la riga precedente se usi questa):
    output_path = os.path.join(
        os.path.dirname(audio_file),
        output_filename
    )

    # Salva il testo trascritto nel file di output
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(result["text"])

    print("Trascrizione completata. File salvato in:", output_path)

    # Genera il report di meeting utilizzando Gemini 2.0 Flash
    generate_report(output_path, os.path.dirname(output_path)) # Salva il report nella stessa cartella della trascrizione

if __name__ == "__main__":
    main()
