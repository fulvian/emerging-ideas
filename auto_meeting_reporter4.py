#!/usr/bin/env python3
import sys
import os
import datetime
import time
import subprocess
import re
import threading
import numpy as np

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

import whisper
import sounddevice as sd
from pydub import AudioSegment
from scipy.signal import butter, filtfilt
import noisereduce as nr  # pip install noisereduce

import google.generativeai as genai

from dotenv import load_dotenv  # pip install python-dotenv

load_dotenv()  # Carica le variabili d'ambiente dal file .env (se presente)

# --- Configurazione ---
MONITOR_INTERVAL = 5  # Secondi tra un controllo e l'altro
MEETING_URLS = ["meet.google.com", "teams.microsoft.com"]

# Cartelle per salvataggi
RECORDING_FOLDER = "/Users/fulvioventura/Library/CloudStorage/GoogleDrive-fulviold@gmail.com/Il mio Drive/Fulviate/professional/registrazioni call"
TRANSCRIPT_FOLDER = "/Users/fulvioventura/Library/CloudStorage/GoogleDrive-fulviold@gmail.com/Il mio Drive/Fulviate/professional/Trascrizioni call"

# Nome file audio: "registrazione" seguito da data e ora
AUDIO_FILENAME_FORMAT = "registrazione_{data}_{ora}.mp3"
SAMPLE_RATE = 44100

# Nome del dispositivo aggregato configurato in Audio MIDI Setup (usa BlackHole 2ch + Microfono)
SYSTEM_AUDIO_DEVICE_NAME = "Dispositivo combinato"

# Moltiplicatori per bilanciare i livelli (regola questi valori in base alle tue prove)
SYSTEM_MULTIPLIER = 0.1     # Per l'audio di sistema
MIC_MULTIPLIER = 0.1        # Per il microfono

# --- Funzioni di Filtro ---
def lowpass_filter(data, cutoff=8000, fs=SAMPLE_RATE, order=4):
    nyq = 0.5 * fs
    normal_cutoff = cutoff / nyq
    b, a = butter(order, normal_cutoff, btype='low', analog=False)
    y = filtfilt(b, a, data)
    return y

# --- Funzione Principale: is_meeting_tab_open ---
def is_meeting_tab_open():
    print("is_meeting_tab_open: Funzione avviata.")
    try:
        script = """
        tell application "Google Chrome"
            repeat with w in windows
                repeat with t in tabs of w
                    set theURL to URL of t
                    if theURL contains "meet.google.com" or theURL contains "teams.microsoft.com" then
                        return true
                    end if
                end repeat
            end repeat
            return false
        end tell
        """
        process = subprocess.run(['osascript', '-e', script], capture_output=True, text=True)
        output = process.stdout.strip().lower()
        result = "true" in output
        print(f"is_meeting_tab_open: Risultato script: {output} -> {result}")
        return result
    except Exception as e:
        print(f"is_meeting_tab_open: Errore: {e}")
        return False

# --- Funzione di registrazione ---
def record_audio_async(stop_event, audio_buffer, input_device, sample_rate):
    def callback(indata, frames, time_info, status):
        if status:
            print(f"record_audio_async: status: {status}")
        audio_buffer.append(indata.copy())
    try:
        with sd.InputStream(callback=callback,
                            samplerate=sample_rate,
                            channels=2,  # Poiché il dispositivo combinato potrebbe avere ingressi multipli
                            device=input_device):
            print("record_audio_async: Registrazione in corso...")
            while not stop_event.is_set():
                time.sleep(0.1)
    except Exception as e:
        print(f"record_audio_async: Errore durante la registrazione: {e}")

# --- Funzioni di Reporting ---
def generate_report_from_audio(audio_file_path, output_path):
    print(f"generate_report_from_audio: Avvio trascrizione per {audio_file_path}")
    try:
        model = whisper.load_model("large")
        result = model.transcribe(audio_file_path)
        base_name = os.path.basename(audio_file_path).replace(".mp3", ".txt")
        transcription_file_path = os.path.join(TRANSCRIPT_FOLDER, base_name)
        with open(transcription_file_path, "w", encoding="utf-8") as f:
            f.write(result["text"])
        print(f"generate_report_from_audio: Trascrizione salvata in {transcription_file_path}")
        generate_report(transcription_file_path, output_path)
    except Exception as e:
        print(f"generate_report_from_audio: Errore durante trascrizione/report: {e}")

def generate_report(transcription_file_path, output_path):
    """
    Genera un report di meeting utilizzando Gemini 2.0 Flash e lo salva in formato DOCX.
    """
    print("Generazione del report di meeting con Gemini 2.0 Flash...")
    try:
        # Configura la chiave API
        genai.configure(api_key=os.environ.get("GOOGLE_API_KEY"))

        # Specifica il modello da utilizzare
        model = genai.GenerativeModel("gemini-2.0-flash")  # oppure "gemini-pro"

        # Leggi il contenuto della trascrizione
        with open(transcription_file_path, "r", encoding="utf-8") as f:
            transcription_text = f.read()

        contents = [
            genai.types.Content(
                role="user",
                parts=[
                    transcription_text,
                    """Sulla base della trascrizione della riunione, crea un resoconto approfondito, completo e dettagliato degli argomenti discussi, senza omettere alcun intervento o argomento, non essere troppo sintetico, non banalizzare i concetti e gli argomenti discussi. Utilizza un linguaggio tecnico e professionale. Individua come prima cosa un titolo dell'incontro sulla base degli argomenti discussi. Individua i partecipanti all'incontro ed elencali in apertura, indicando anche la data dell'incontro che puoi ricavare dal titolo del file. Struttura il report in paragrafi suddividendolo per tematiche omogenee. Non citare mai direttamente gli speaker. Al termine del resoconto evidenzia gli appuntamenti futuri, le azioni da mettere in atto e a chi si riferiscono. All'inizio del documento crea un executive summary di quanto discusso, delle decisioni prese e delle cose da fare, in forma di punto elenco. Restituisci il testo in italiano senza eccedere in sinteticità, è fondamentale la completezza delle informazioni desumibili."""
                ],
            )
        ]

        generation_config = genai.types.GenerationConfig(
            temperature=0.2,
            top_p=1.00,
            top_k=40,
            max_output_tokens=60000,
            response_mime_type="text/plain",
        )
        report_text = ""
        response = model.generate_content(
            contents=contents,
            generation_config=generation_config,
            stream=True,
        )
        for chunk in response:
            report_text += chunk.text

        # Creazione del documento DOCX
        document = Document()
        styles = document.styles
        for style in styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                style.font.name = 'Calibri Light'
                style._element.rPr.rFonts.set(qn('w:asciiTheme'), 'Calibri Light')
                style._element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Calibri Light')

        report_title = None
        for line in report_text.splitlines():
            line = line.strip()
            if line:
                report_title = line.lstrip('#').strip()
                document.add_heading(report_title, level=1)
                break

        if report_title:
            report_title_cleaned = re.sub(r'[^\w\s-]', '', report_title).strip().replace(' ', '_')
            today_str = datetime.datetime.now().strftime("%d_%m_%Y")
            report_filename = f"{report_title_cleaned}_{today_str}.docx"
        else:
            today_str = datetime.datetime.now().strftime("%d_%m_%Y")
            report_filename = f"report_senza_titolo_{today_str}.docx"

        report_filepath = os.path.join(output_path, report_filename)
        for line in report_text.splitlines():
            line = line.strip()
            if line:
                document.add_paragraph(line.lstrip('#').strip())
        document.save(report_filepath)
        print("Report di meeting generato e salvato in:", report_filepath)
    except Exception as e:
        print(f"Si è verificato un errore durante la generazione del report: {e}")

# --- Funzione di Monitoraggio ---
def monitor_and_record():
    print("monitor_and_record: Avvio monitoraggio.")
    is_recording = False
    audio_filename = None
    recording_thread = None
    stop_event = None
    audio_buffer =# Inizializzazione corretta di audio_buffer

    # Ottieni l'indice del dispositivo usando sounddevice
    pa_devices = sd.query_devices()
    device_index = None
    for i, dev in enumerate(pa_devices):
        if SYSTEM_AUDIO_DEVICE_NAME.lower() in dev['name'].lower():
            device_index = i
            break
    if device_index is None:
        print(f"monitor_and_record: Dispositivo '{SYSTEM_AUDIO_DEVICE_NAME}' non trovato. Esco.")
        sys.exit(1)
    else:
        print(f"monitor_and_record: Trovato dispositivo '{SYSTEM_AUDIO_DEVICE_NAME}' con indice {device_index}.")

    while True:
        print("monitor_and_record: Verifica scheda meeting...")
        if is_meeting_tab_open():
            print("monitor_and_record: Meeting rilevato.")
            if not is_recording:
                is_recording = True
                now = datetime.datetime.now()
                date_str = now.strftime("%d_%m_%Y")
                time_str = now.strftime("%H_%M")
                audio_filename = os.path.join(RECORDING_FOLDER, AUDIO_FILENAME_FORMAT.format(data=date_str, ora=time_str))
                print(f"monitor_and_record: Inizio registrazione su {SYSTEM_AUDIO_DEVICE_NAME} (indice {device_index})...")
                stop_event = threading.Event()
                audio_buffer =# Reinizializza buffer prima del thread
                recording_thread = threading.Thread(
                    target=record_audio_async,
                    args=(stop_event, audio_buffer, SYSTEM_AUDIO_DEVICE_NAME, SAMPLE_RATE)
                )
                recording_thread.start()
            else:
                print("monitor_and_record: Registrazione in corso...")
            time.sleep(MONITOR_INTERVAL)
        else:
            print("monitor_and_record: Meeting non rilevato.")
            if is_recording:
                print("monitor_and_record: Terminazione registrazione...")
                stop_event.set()
                recording_thread.join()
                is_recording = False
                if audio_buffer and len(audio_buffer) > 0:
                    recorded = np.concatenate(audio_buffer, axis=0)
                    if recorded.ndim == 2 and recorded.shape[1] >= 3:
                        # Se abbiamo 3 canali, assumiamo: canale 0 = system, canale 2 = microfono
                        system_channel = recorded[:, 0] * SYSTEM_MULTIPLIER
                        mic_channel = recorded[:, 2] * MIC_MULTIPLIER
                        mixed = np.column_stack((system_channel, mic_channel))
                        mixed[:, 0] = lowpass_filter(mixed[:, 0])
                        mixed[:, 0] = nr.reduce_noise(y=mixed[:, 0], sr=SAMPLE_RATE)
                        mixed[:, 1] = lowpass_filter(mixed[:, 1])
                        peak = np.max(np.abs(mixed))
                        if peak > 0:
                            scaling_factor = 0.9 / peak
                            mixed *= scaling_factor
                        mixed = np.clip(mixed, -1.0, 1.0)
                    elif recorded.ndim == 2 and recorded.shape[1] == 2:
                        system_channel = recorded[:, 0] * SYSTEM_MULTIPLIER
                        mic_channel = recorded[:, 1] * MIC_MULTIPLIER
                        mixed = np.column_stack((system_channel, mic_channel))
                        mixed[:, 0] = lowpass_filter(mixed[:, 0])
                        mixed[:, 1] = lowpass_filter(mixed[:, 1])
                        mixed[:, 0] = nr.reduce_noise(y=mixed[:, 0], sr=SAMPLE_RATE)
                        mixed[:, 1] = nr.reduce_noise(y=mixed[:, 1], sr=SAMPLE_RATE)
                        peak = np.max(np.abs(mixed))
                        if peak > 0:
                            scaling_factor = 0.9 / peak
                            mixed *= scaling_factor
                        mixed = np.clip(mixed, -1.0, 1.0)
                    else:
                        mixed = recorded
                    print("monitor_and_record: Salvataggio file audio...")
                    audio_segment = AudioSegment(
                        mixed.tobytes(),
                        frame_rate=SAMPLE_RATE,
                        sample_width=mixed.dtype.itemsize,
                        channels=2
                    )
                    audio_segment.export(audio_filename, format="mp3")
                    print(f"monitor_and_record: Registrazione salvata in: {audio_filename}")
                    output_path = os.path.dirname(audio_filename)
                    generate_report_from_audio(audio_filename, output_path)
                else:
                    print("monitor_and_record: Nessun audio registrato.")
        print(f"monitor_and_record: Attesa {MONITOR_INTERVAL} secondi...")
        time.sleep(MONITOR_INTERVAL)

if __name__ == "__main__":
    print("Script di monitoraggio avviato in background...")
    monitor_and_record()
